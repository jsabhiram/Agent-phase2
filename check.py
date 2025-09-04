from docx import Document
import re
import os

def parse_invoice_items(path):
    """
    Extracts purchase item details from a DOCX invoice.

    Args:
        path (str): Path to the DOCX invoice file.

    Returns:
        list[dict]: Each item includes:
            - name (str)
            - quantity (str)
            - price (str)
            - total (str)
    """
    doc = Document(path)
    items = []

    # Regex patterns
    item_pattern = re.compile(r'(?P<quantity>\d+)\s+(?P<name>.+?)\s+(?P<price>[\d,]+\.\d{2})\s+(?P<total>[\d,]+\.\d{2})')
    discount_pattern = re.compile(r'discount|disc|deduction|less', re.IGNORECASE)
    numeric_pattern = re.compile(r'[\d,]+\.\d{2}')

    # Attempt table parsing
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]

            if len(cells) < 3 or any(x.lower() in ['item', 'description', 'qty'] for x in cells):
                continue

            try:
                if len(cells) >= 4:
                    item = {
                        'name': cells[1],
                        'quantity': cells[2],
                        'price': cells[3].replace(',', ''),
                        'total': cells[-1].replace(',', '')
                    }
                elif len(cells) == 3:
                    item = {
                        'name': cells[0],
                        'quantity': '1',
                        'price': cells[1].replace(',', ''),
                        'total': cells[2].replace(',', '')
                    }
                else:
                    continue

                items.append(item)
            except Exception:
                continue

    # Fallback: Paragraph-based parsing
    if not items:
        for para in doc.paragraphs:
            text = para.text.strip()

            if not text or text.lower().startswith(('invoice', 'date', 'page')):
                continue

            match = item_pattern.search(text)
            if match:
                items.append(match.groupdict())
                continue

            if discount_pattern.search(text):
                nums = numeric_pattern.findall(text)
                if nums:
                    items.append({
                        'name': 'DISCOUNT',
                        'quantity': '1',
                        'price': f"-{nums[0].replace(',', '')}",
                        'total': f"-{nums[-1].replace(',', '')}"
                    })

    # Remove unwanted items (e.g., Shipping or empty names)
    items = [x for x in items if not (x.get('name', '').lower().startswith('shipping') or x.get('name', '').strip() == '')]

    return items

# Optional: Main method for testing or CLI use
def main(file_path):
    path =file_path
    invoice_items = parse_invoice_items(path)
    return invoice_items



if __name__ == "__main__":
    print("Module loaded")
    # for item in invoice_items:
    # print(f"Item: {item.get('name')}")
    # print(f"Quantity: {item.get('quantity')}")
    # print(f"Price: {item.get('price')}")
    # print(f"Total: {item.get('total')}")
    # print("-" * 30)
